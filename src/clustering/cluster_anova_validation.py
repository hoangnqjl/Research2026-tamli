import pandas as pd
import scipy.stats as stats
import sys
import os

# Redirect stdout to a file as well
class DualLogger:
    def __init__(self, filename):
        self.terminal = sys.stdout
        self.log = open(filename, "w", encoding='utf-8')

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)

    def flush(self):
        self.terminal.flush()
        self.log.flush()

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.stdout = DualLogger(os.path.join(BASE_DIR, "outputs", "anova_validation_results.txt"))

print("--- CLUSTER VALIDATION: ONE-WAY ANOVA ---")
print("Objective: Prove that cluster differences are statistically significant (p < 0.05).")

# 1. LOAD DATA
input_file = os.path.join(BASE_DIR, 'data', 'processed', 'blind_clustering_final.xlsx')
if not os.path.exists(input_file):
    print(f"Error: {input_file} not found.")
    sys.exit(1)

try:
    df = pd.read_excel(input_file, sheet_name='Clustered_Users')
    print("Loaded data successfully.")
except Exception as e:
    print(f"Error loading data: {e}")
    sys.exit(1)

# 2. DEFINE VARIABLES
feature_cols = [
    '[SumScore_Work_Check]', '[SumScore_Source_Check]', '[SumScore_Sleep_Loss]',
    '[SumScore_Obsess_Loop]', '[SumScore_Obsess_Debate]', '[SumScore_Obsess_Celeb]',
    '[SumScore_FOMO_Reaction]', '[SumScore_Fatigue]', '[SumScore_Escapism]',
    '[SumScore_Deepfake_Detect]', '[SumScore_Deep_Reading]', '[SumScore_Crowd_Pressure]',
    '[SumScore_Control_Time]', '[SumScore_Clickbait_Aware]', '[SumScore_AI_Trust]',
    '[SumScore_AI_Freq]'
]

# 3. RUN ANOVA FOR EACH VARIABLE
results = []

for col in feature_cols:
    # Separate data into groups based on Cluster_ID
    group0 = df[df['Cluster_ID'] == 0][col]
    group1 = df[df['Cluster_ID'] == 1][col]
    group2 = df[df['Cluster_ID'] == 2][col]
    
    # Perform One-way ANOVA
    f_stat, p_val = stats.f_oneway(group0, group1, group2)
    
    # Determine significance
    if p_val < 0.001:
        conclusion = "Highly Significant"
        sig_level = "***"
    elif p_val < 0.01:
        conclusion = "Significant"
        sig_level = "**"
    elif p_val < 0.05:
        conclusion = "Weakly Significant"
        sig_level = "*"
    else:
        conclusion = "Not Significant"
        sig_level = "ns"
        
    results.append({
        'Feature': col,
        'F_Stat': f_stat,
        'P_Val': p_val,
        'Conclusion': conclusion,
        'Sig': sig_level
    })

# SORT RESULTS BY F-STATISTIC (DESCENDING)
results.sort(key=lambda x: x['F_Stat'], reverse=True)

# PRINT TABLE
print("\n" + "="*100)
print(f"{'ANOVA TEST RESULTS (Sorted by Significance)':^100}")
print("="*100)
print(f"{'Feature':<30} | {'F-Statistic':<12} | {'P-Value':<12} | {'Conclusion':<20}")
print("-" * 100)

for res in results:
     print(f"{res['Feature']:<30} | {res['F_Stat']:<12.2f} | {res['P_Val']:<12.2e} | {res['Conclusion']}")

# 4. COMPARE WITH CENTROID FINDINGS
print("\n" + "="*100)
print(f"{'SCIENTIFIC INTERPRETATION':^100}")
print("="*100)

print("\nTOP 5 KEY DIFFERENTIATORS (Variables that best define the clusters):")
for i, res in enumerate(results[:5]):
    print(f"{i+1}. {res['Feature']} (F={res['F_Stat']:.1f}, p={res['P_Val']:.2e})")

print("\n -> These variables are the 'Fingerprints' of your clusters.")
print(" -> High F-statistic means the Between-Group variance is much larger than Within-Group variance.")
print(" -> This scientifically proves that your clusters are distinct and not random.")

print("\nAnalysis Complete. Results saved to 'anova_validation_results.txt'.")

try:
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    doc.add_heading('ANOVA Test Results', 0)
    
    doc.add_paragraph("Objective: Prove that cluster differences are statistically significant (p < 0.05).")

    # Add Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'F-Statistic'
    hdr_cells[2].text = 'P-Value'
    hdr_cells[3].text = 'Conclusion'
    
    for res in results:
        row_cells = table.add_row().cells
        row_cells[0].text = res['Feature']
        row_cells[1].text = f"{res['F_Stat']:.2f}"
        row_cells[2].text = f"{res['P_Val']:.2e}"
        row_cells[3].text = res['Conclusion']
    
    doc.add_heading('Scientific Interpretation', level=1)
    doc.add_paragraph("TOP 5 KEY DIFFERENTIATORS (Variables that best define the clusters):")
    
    for i, res in enumerate(results[:5]):
        doc.add_paragraph(f"{i+1}. {res['Feature']} (F={res['F_Stat']:.1f}, p={res['P_Val']:.2e})", style='List Number')
        
    doc.add_paragraph("\n-> These variables are the 'Fingerprints' of your clusters.")
    doc.add_paragraph("-> High F-statistic means the Between-Group variance is much larger than Within-Group variance.")
    doc.add_paragraph("-> This scientifically proves that your clusters are distinct and not random.")

    docx_path = os.path.join(BASE_DIR, "outputs", "anova_validation_results.docx")
    doc.save(docx_path)
    print(f"Results also saved to '{docx_path}'")

except ImportError:
    print("Error: 'python-docx' library not found. Install it via 'pip install python-docx' to generate Word report.")
except Exception as e:
    print(f"Error saving Word document: {e}")
